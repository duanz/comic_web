import mkepub
import os
import shutil
import logging
from comic_admin.models import Comic, ChapterImage, CoverImage
from comic_admin.models import Chapter as ComicChapter
from book_admin.models import Book, Author
from book_admin.models import Chapter as BookChapter
from comic_admin.serializers import CommonImageSerializer
from django.conf import settings
from docx import Document
from docx.shared import Inches
from comic_web.utils.photo import split_photo_fit_kindle

class MakeMyEpub:
    def __init__(self, content_id, content_type):
        self.content_id = content_id
        self.content_type = content_type
        self.title = ""
        self.filename = ""

    def run(self):
        if self.content_type == "comic":
            self.makeComicWord()
        elif self.content_type == 'book':
            # self.makeBookEpub()
            self.makeBookTxt()
        pass

    def makeComicEpub(self):
        comic_id = self.content_id
        comic_obj = Comic.normal.filter(id=comic_id).first()
        self.title = comic_obj.title

        # 初始化epub
        epub = mkepub.Book(title=comic_obj.title, author=str(comic_obj.author))
        logging.info("epub 已经初始化")

        # 设置封面
        comic_cover_path_list = CommonImageSerializer.to_representation(comic_id=comic_id, img_type="comic_cover", quality="title", only_url=False, only_path=True)
        logging.info(comic_cover_path_list)
        if comic_cover_path_list:
            with open(comic_cover_path_list[0], 'rb') as file:
                epub.set_cover(file.read())
        logging.info("epub 设置封面完成")

        # 设置章节
        chapters = ComicChapter.normal.filter(comic_id=comic_id)
        for chapter in chapters:
            html = ""
            img_lab = '<img src="images/{}" width="100%" alt="You can use images as well">'
            chapter_img_path_list = CommonImageSerializer.to_representation(chapter_id=chapter.id, img_type="chapter_content", quality="thumb", only_url=False, only_path=True)
            if chapter_img_path_list:
                for img_idx, img_path in enumerate(chapter_img_path_list):
                    file_name = "{}_{}_{}.jpg".format(comic_id, chapter.id, img_idx)
                    html+=img_lab.format(file_name)
                    with open(img_path, 'rb') as file:
                        epub.add_image(file_name, file.read())
            epub.add_page(chapter.title, html)
        
        logging.info("epub 完成")
        filename = os.path.join(settings.UPLOAD_SAVE_PATH, comic_obj.title + '.epub')
        if os.path.exists(filename):
            os.remove(filename)
        self.filename = filename
        epub.save(filename)
        comic_obj.is_download = True
        comic_obj.save()

    def makeBookEpub(self):
        book_id = self.content_id
        book_obj = Book.normal.filter(id=book_id).first()
        self.title = book_obj.title

        # 初始化epub
        epub = mkepub.Book(title=book_obj.title, author=str(book_obj.author))

        # 设置封面
        book_cover_path_list = CommonImageSerializer.to_representation(book_id=book_id, img_type="book_cover", only_url=False, only_path=True)
        if book_cover_path_list:
            with open(book_cover_path_list[0], 'rb') as file:
                epub.set_cover(file.read())

        # 设置章节
        chapters = BookChapter.normal.filter(book_id=book_id)
        for chapter in chapters:
            epub.add_page(chapter.title, chapter.content)
        
        filename = os.path.join(settings.UPLOAD_SAVE_PATH, book_obj.title + '.epub')
        self.filename = filename
        epub.save(filename)
        book_obj.is_download = True
        book_obj.save()

    def makeBookTxt(self):
        book_id = self.content_id
        book_obj = Book.normal.filter(id=book_id).first()
        self.title = book_obj.title

        # 初始化txt
        if not os.path.exists(settings.UPLOAD_SAVE_PATH):
            os.makedirs(settings.UPLOAD_SAVE_PATH, 0o775)

        filename = os.path.join(settings.UPLOAD_SAVE_PATH, book_obj.title + '.txt')
        if os.path.exists(filename):
            os.remove(filename)
        self.filename = filename
        book = open(filename, 'w+')        

        # 设置章节
        chapters = BookChapter.normal.filter(book_id=book_id)
        for chapter in chapters:
            book.writelines(chapter.title+'\n')
            book.writelines(chapter.content+'\n')
        
        book.flush()
        book.close()
        
        book_obj.is_download = True
        book_obj.save()

    def makeComicWord(self):
        comic_id = self.content_id
        comic_obj = Comic.normal.filter(id=comic_id).first()
        self.title = comic_obj.title

        # 临时文件夹
        comic_temp_path = os.path.join(settings.UPLOAD_SAVE_PATH, self.title)

        part = 0
        part_size = 1024 * 1024 * 20
        current_size = 0
        pre_size = lambda cur: cur + 1024*1024*5
        # 设置章节
        chapters = ComicChapter.normal.filter(comic_id=comic_id)
        for chapter in chapters:
            if current_size == 0:
                # 初始化word
                part += 1
                doc = Document()
                doc.add_heading(chapter.title, level=1)
                logging.info("WORD part-{} 已经初始化".format(part))
            

            chapter_img_path_list = CommonImageSerializer.to_representation(chapter_id=chapter.id, img_type="chapter_content", quality="title", only_url=False, only_path=True)
            if chapter_img_path_list:
                for img_idx, img_path in enumerate(chapter_img_path_list):
                    img_size = os.path.getsize(img_path)
                    current_size += img_size

                    # 切割大图片临时文件夹
                    temp_path = os.path.join(comic_temp_path, os.path.split(img_path)[-1].split('.')[0])

                    # 如果是大文件就分隔
                    after_split = split_photo_fit_kindle(img_path, temp_path)
                    for small_img in after_split:
                        doc.add_picture(small_img)

            if pre_size(current_size) >= part_size:
                # 保存word
                filename = os.path.join(settings.UPLOAD_SAVE_PATH, '{}__{}.docx'.format(comic_obj.title, part))
                if os.path.exists(filename):
                    os.remove(filename)
                doc.save(filename)
                current_size = 0
                logging.info("WORD part-{} 已经完成".format(part))
        
        # 删除临时文件
        shutil.rmtree(comic_temp_path)

        logging.info("word 完成")
        comic_obj.is_download = True
        comic_obj.save()


# class FileOperationBase:
#     '''分割文件为20M/个'''
# 	def __init__(self, filename, srcpath, despath, chunksize=1024*20):
#         self.filename = filename
#         self.chunksize = chunksize
# 		self.srcpath = srcpath
# 		self.despath = despath
 
# 	def splitFile(self):
# 		'split the files into chunks, and save them into despath'
# 		if not os.path.exists(self.despath):
# 			os.mkdir(self.despath)
# 		chunknum = 0
# 		inputfile = open(self.srcpath, 'rb') #rb 读二进制文件
# 		try:
# 			while 1:
# 				chunk = inputfile.read(self.chunksize)
# 				if not chunk: #文件块是空的
# 					break
# 				chunknum += 1
# 				filename = os.path.join(self.despath, ("{}-part{}".format(self.filename, chunknum)))
# 				fileobj = open(filename, 'wb')
# 				fileobj.write(chunk)
# 		except IOError:
# 			print "read file error\n"
# 			raise IOError
# 		finally:
# 			inputfile.close()
# 		return chunknum